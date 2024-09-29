from datetime import datetime, timedelta
import matplotlib.pyplot as plt
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem,Image
from reportlab.lib.styles import getSampleStyleSheet
from textblob import TextBlob
from docx import Document
from datetime import datetime, timedelta
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Table, TableStyle
import docx
from docx.shared import Inches
import re
from collections import Counter

# Function to categorize sentiment based on polarity score
def categorize_sentiment(polarity):
    if polarity > 0.2:
        return "Positive"
    elif polarity < -0.2:
        return "Negative"
    else:
        return "Neutral"

# Function to analyze speech and categorize sentiment
def analyze_speech(transcript_data):
    analysis_results = []
    sentiment_summary = {"Positive": 0, "Neutral": 0, "Negative": 0}

    for entry in transcript_data:
        speaker = entry['name']
        content = entry['content']

        # Perform sentiment analysis
        blob = TextBlob(content)
        sentiment = blob.sentiment
        sentiment_category = categorize_sentiment(sentiment.polarity)

        # Increment sentiment summary count
        sentiment_summary[sentiment_category] += 1

        results = {
            'speaker': speaker,
            'content': content,
            'polarity': sentiment.polarity,
            'subjectivity': sentiment.subjectivity,
            'sentiment_category': sentiment_category
        }
        analysis_results.append(results)
    
    return analysis_results, sentiment_summary

# Function to generate overall summary and key takeaways based on analysis
def generate_summary_and_takeaways(analysis_results):
    overall_summary = "Summary of the meeting"
    key_takeaways = []

    if analysis_results:
        for entry in analysis_results:
            key_takeaways.append(f"{entry['speaker']} mentioned: '{entry['content']}' with sentiment polarity of {entry['polarity']:.2f}.")
    
    return overall_summary, key_takeaways

# Function to create a single report with interval sections
def create_report_with_interval_sections(meeting_data, interval_minutes):
    file_name = "Complete_Interval_Report.pdf"
    pdf = SimpleDocTemplate(file_name, pagesize=A4)
    elements = []
    
    styles = getSampleStyleSheet()
    title = f"<b>Meeting Report: {meeting_data['meetingTitle']}</b>"
    elements.append(Paragraph(title, styles['Title']))
    elements.append(Spacer(1, 12))

    start_time = datetime.fromisoformat(meeting_data['meetingStartTimeStamp'].replace("Z", "+00:00"))
    end_time = datetime.fromisoformat(meeting_data['meetingEndTimeStamp'].replace("Z", "+00:00"))
    
    interval_seconds = interval_minutes * 60
    current_time = start_time

    # Process intervals and add sections to the report
    while current_time < end_time:
        interval_end_time = min(current_time + timedelta(seconds=interval_seconds), end_time)

        # Filter transcript data for the current interval
        interval_transcript_data = [
            entry for entry in meeting_data['transcriptData']
            if current_time <= datetime.fromisoformat(entry['timestamp'].replace("Z", "+00:00")) < interval_end_time
        ]
        
        if interval_transcript_data:
            # Create section title for the interval
            interval_title = f"<b>Interval: {current_time.strftime('%H:%M')} - {interval_end_time.strftime('%H:%M')}</b>"
            elements.append(Paragraph(interval_title, styles['Heading2']))
            elements.append(Spacer(1, 12))
            
            # Analyze speech for this interval
            analysis_results, sentiment_summary = analyze_speech(interval_transcript_data)

            # Generate overall summary and key takeaways
            overall_summary, key_takeaways = generate_summary_and_takeaways(analysis_results)

            # Add overall summary for this interval
            elements.append(Paragraph("<b>Summary:</b>", styles['Heading3']))
            elements.append(Paragraph(overall_summary, styles['Normal']))
            elements.append(Spacer(1, 12))

            # Add key takeaways as bullet points
            elements.append(Paragraph("<b>Key Takeaways:</b>", styles['Heading3']))
            bullet_points = [ListItem(Paragraph(f"• {takeaway}", styles['Normal']), leftIndent=10) for takeaway in key_takeaways]
            elements.append(ListFlowable(bullet_points, bulletType='bullet'))
            elements.append(Spacer(1, 12))

            # Add sentiment summary for this interval
            elements.append(Paragraph("<b>Sentiment Summary:</b>", styles['Heading3']))
            for sentiment, count in sentiment_summary.items():
                elements.append(Paragraph(f"{sentiment}: {count} occurrences", styles['Normal']))
                elements.append(Spacer(1, 12))
        
        current_time = interval_end_time

    # Build the final PDF
    pdf.build(elements)
    print(f"{file_name} generated successfully.")


# Helper function to format timestamp to HH:MM:SS and handle the fractional seconds and 'Z'
def format_time(timestamp):
    timestamp = timestamp.split('.')[0]
    return datetime.strptime(timestamp, '%Y-%m-%dT%H:%M:%S').strftime('%H:%M:%S')

# Function to dynamically generate the overall summary based on the transcript
def generate_overall_summary(transcript_data):
    content_list = [entry['content'].strip() for entry in transcript_data]
    all_words = re.findall(r'\b\w{4,}\b', ' '.join(content_list).lower())
    common_words = Counter(all_words).most_common(3)

    if common_words:
        summary = f"In this meeting, the main discussions focused on: {', '.join([word for word, _ in common_words])}."
    else:
        summary = "General discussions took place in this meeting."

    return summary

# Function to dynamically generate key takeaways based on the transcript
def generate_key_takeaways(transcript_data):
    takeaways = []

    for entry in transcript_data:
        content = entry['content'].strip().lower()
        if "question" in content:
            takeaways.append(f"{entry['name']} asked a question.")
        if "discussed" in content:
            takeaways.append(f"{entry['name']} discussed {content}.")

    if not takeaways:
        takeaways.append("General discussion, no specific key takeaways.")

    return takeaways

# Function to dynamically generate FAQs based on the transcript
def generate_faqs(transcript_data):
    faqs = []
    
    for entry in transcript_data:
        content = entry['content'].strip().lower()
        if "?" in content:  # Check for questions in the content
            faqs.append(f"Q: {content}")
        elif "answer" in content:  # Assuming answers follow the question
            faqs.append(f"A: {content}")

    if not faqs:
        faqs.append("No FAQs were raised during the meeting.")

    return faqs

# Function to generate speaker summaries
def generate_speaker_summaries(transcript_data, speaker_durations):
    speaker_summaries = {}

    for entry in transcript_data:
        speaker = entry['name']
        content = entry['content'].strip()

        if speaker not in speaker_summaries:
            speaker_summaries[speaker] = {
                "summary": f"{speaker} contributed the following points:\n",
                "points": [],
                "duration": speaker_durations.get(speaker, 0)
            }
        speaker_summaries[speaker]['points'].append(f"Said: \"{content}\"")

    return speaker_summaries

# Function to create the normal report PDF
def create_normal_report_pdf(meeting_data):
    file_name = f"{meeting_data['meetingTitle']}_summary_report.pdf"
    pdf = SimpleDocTemplate(file_name, pagesize=A4)
    elements = []

    styles = getSampleStyleSheet()
    title = f"<b>{meeting_data['meetingTitle']}</b>"
    elements.append(Paragraph(title, styles['Title']))
    elements.append(Spacer(1, 12))

    # Add meeting details with formatted timestamps
    meeting_details = (
        f"Meeting Start Time: {format_time(meeting_data.get('meetingStartTimeStamp', ''))}<br/>"
        f"Meeting End Time: {format_time(meeting_data.get('meetingEndTimeStamp', ''))}<br/>"
        f"Attendees: {', '.join(meeting_data.get('attendees', ['N/A']))}<br/><br/>"
    )
    elements.append(Paragraph(meeting_details, styles['Normal']))

    # Generate and add Executive Summary
    overall_summary = generate_overall_summary(meeting_data.get('transcriptData', []))
    elements.append(Paragraph("<b>Executive Summary:</b>", styles['Heading2']))
    elements.append(Paragraph(overall_summary, styles['Normal']))
    elements.append(Spacer(1, 12))

    # Generate and add Key Takeaways
    key_takeaways = generate_key_takeaways(meeting_data.get('transcriptData', []))
    elements.append(Paragraph("<b>Key Takeaways:</b>", styles['Heading2']))
    bullet_points = [ListItem(Paragraph(f"• {takeaway}", styles['Normal']), leftIndent=10) for takeaway in key_takeaways]
    elements.append(ListFlowable(bullet_points, bulletType='bullet'))
    elements.append(Spacer(1, 12))

    # Generate and add Speaker Summaries
    speaker_summaries = generate_speaker_summaries(meeting_data.get('transcriptData', []), meeting_data.get('speakerDuration', {}))
    elements.append(Paragraph("<b>Speaker Summaries:</b>", styles['Heading2']))
    for speaker, data in speaker_summaries.items():
        speaker_info = f"<b>{speaker}</b><br/>Total Speaking Time: {data['duration']} seconds<br/>"
        elements.append(Paragraph(speaker_info, styles['Normal']))
        for point in data['points']:
            elements.append(Paragraph(f"• {point}", styles['Normal']))
        elements.append(Spacer(1, 12))

    # Add FAQ Section
    elements.append(Paragraph("<b>Frequently Asked Questions (FAQ):</b>", styles['Heading2']))
    faqs = generate_faqs(meeting_data.get('transcriptData', []))
    for faq in faqs:
        elements.append(Paragraph(f"• {faq}", styles['Normal']))
    elements.append(Spacer(1, 12))

    # Add Appendix Section
    elements.append(Paragraph("<b>Appendix:</b>", styles['Heading2']))
    appendix_content = (
        "This report includes detailed discussions on project timelines, roles, and responsibilities. "
        "For more information, refer to the shared documents."
    )
    elements.append(Paragraph(appendix_content, styles['Normal']))
    
    # Finalize and build the PDF
    pdf.build(elements)
    print(f"PDF report '{file_name}' generated successfully.")

# Function to create the normal report DOCX
def create_normal_report_docx(meeting_data):
    file_name = f"{meeting_data['meetingTitle']}_summary_report.docx"
    doc = Document()

    doc.add_heading(meeting_data['meetingTitle'], level=1)

    # Add meeting details
    doc.add_paragraph(
        f"Meeting Start Time: {format_time(meeting_data.get('meetingStartTimeStamp', ''))}\n"
        f"Meeting End Time: {format_time(meeting_data.get('meetingEndTimeStamp', ''))}\n"
        f"Attendees: {', '.join(meeting_data.get('attendees', ['N/A']))}\n"
    )

    # Generate and add Executive Summary
    overall_summary = generate_overall_summary(meeting_data.get('transcriptData', []))
    doc.add_heading('Executive Summary:', level=2)
    doc.add_paragraph(overall_summary)

    # Generate and add Key Takeaways
    key_takeaways = generate_key_takeaways(meeting_data.get('transcriptData', []))
    doc.add_heading('Key Takeaways:', level=2)
    for takeaway in key_takeaways:
        doc.add_paragraph(f"• {takeaway}")

    # Generate and add Speaker Summaries
    speaker_summaries = generate_speaker_summaries(meeting_data.get('transcriptData', []), meeting_data.get('speakerDuration', {}))
    doc.add_heading('Speaker Summaries:', level=2)
    for speaker, data in speaker_summaries.items():
        doc.add_paragraph(f"{speaker} - Total Speaking Time: {data['duration']} seconds")
        for point in data['points']:
            doc.add_paragraph(f"• {point}")

    # Add FAQ Section
    doc.add_heading('Frequently Asked Questions (FAQ):', level=2)
    faqs = generate_faqs(meeting_data.get('transcriptData', []))
    for faq in faqs:
        doc.add_paragraph(f"• {faq}")

    # Add Appendix Section
    doc.add_heading('Appendix:', level=2)
    appendix_content = (
        "This report includes detailed discussions on project timelines, roles, and responsibilities. "
        "For more information, refer to the shared documents."
    )
    doc.add_paragraph(appendix_content)

    # Save the DOCX file
    doc.save(file_name)
    print(f"DOCX report '{file_name}' generated successfully.")

def create_speaker_ranking_report_pdf(meeting_data):
    file_name = "Speaker_Ranking_Report.pdf"
    pdf = SimpleDocTemplate(file_name, pagesize=A4)
    elements = []
    
    styles = getSampleStyleSheet()
    title = f"<b>Speaker Ranking Report: {meeting_data['meetingTitle']}</b>"
    elements.append(Paragraph(title, styles['Title']))
    elements.append(Spacer(1, 12))

    # Prepare to collect speaker data
    speaker_data = {}

    # Process transcript data to gather speaker information
    for entry in meeting_data['transcriptData']:
        speaker_name = entry['name']
        content = entry['content']
        timestamp = datetime.fromisoformat(entry['timestamp'].replace("Z", "+00:00"))

        if speaker_name not in speaker_data:
            speaker_data[speaker_name] = {
                'count': 0,
                'duration': 0,
                'summary': []  # List to hold the speaker's speech summary
            }
        
        speaker_data[speaker_name]['count'] += 1
        speaker_data[speaker_name]['summary'].append(content)  # Collect each speech content
        speaker_data[speaker_name]['duration'] += 1  # Assuming each entry has a duration of 1 minute for simplicity

    # Sort speakers by their count (descending)
    sorted_speakers = sorted(speaker_data.items(), key=lambda item: item[1]['count'], reverse=True)

    # Add speaker ranking to the PDF
    elements.append(Paragraph("<b>Speaker Ranking:</b>", styles['Heading2']))
    for rank, (speaker_name, data) in enumerate(sorted_speakers, start=1):
        elements.append(Paragraph(f"{rank}. {speaker_name} - {data['count']} contributions, Duration: {data['duration']} minutes", styles['Normal']))
        
        # Add summary for each speaker
        elements.append(Paragraph("<b>Summary of Contributions:</b>", styles['Heading3']))
        for summary in data['summary']:
            elements.append(Paragraph(summary, styles['Normal']))
            elements.append(Spacer(1, 6))  # Space between summaries

    # Build the final PDF
    pdf.build(elements)
    print(f"{file_name} generated successfully.")
# Sample implementation of a speaker ranking report in DOCX
def create_speaker_ranking_report_docx(meeting_data):
    file_name = "Speaker_Ranking_Report.docx"
    doc = Document()

    doc.add_heading(f'Speaker Ranking Report: {meeting_data["meetingTitle"]}', level=1)

    # Prepare to collect speaker data
    speaker_data = {}

    # Process transcript data to gather speaker information
    for entry in meeting_data['transcriptData']:
        speaker_name = entry['name']
        content = entry['content']
        
        if speaker_name not in speaker_data:
            speaker_data[speaker_name] = {
                'count': 0,
                'summary': []  # List to hold the speaker's speech summary
            }
        
        speaker_data[speaker_name]['count'] += 1
        speaker_data[speaker_name]['summary'].append(content)  # Collect each speech content

    # Sort speakers by their count (descending)
    sorted_speakers = sorted(speaker_data.items(), key=lambda item: item[1]['count'], reverse=True)

    # Add speaker ranking to the document
    doc.add_heading('Speaker Ranking:', level=2)
    for rank, (speaker_name, data) in enumerate(sorted_speakers, start=1):
        doc.add_paragraph(f"{rank}. {speaker_name} - {data['count']} contributions")

        # Add summary for each speaker
        doc.add_heading("Summary of Contributions:", level=3)
        for summary in data['summary']:
            doc.add_paragraph(summary)

    # Save the DOCX file
    doc.save(file_name)
    print(f"{file_name} generated successfully.")

# Function to categorize sentiment based on polarity score
def categorize_sentiment(polarity):
    if polarity > 0.2:
        return "Positive"
    elif polarity < -0.2:
        return "Negative"
    else:
        return "Neutral"

# Function to analyze speech and categorize sentiment
def analyze_speech(transcript_data):
    analysis_results = []
    sentiment_summary = {"Positive": 0, "Neutral": 0, "Negative": 0}

    for entry in transcript_data:
        speaker = entry['name']
        content = entry['content']

        # Perform sentiment analysis
        blob = TextBlob(content)
        sentiment = blob.sentiment
        sentiment_category = categorize_sentiment(sentiment.polarity)

        # Increment sentiment summary count
        sentiment_summary[sentiment_category] += 1

        results = {
            'speaker': speaker,
            'content': content,
            'polarity': sentiment.polarity,
            'subjectivity': sentiment.subjectivity,
            'sentiment_category': sentiment_category
        }
        analysis_results.append(results)
    
    return analysis_results, sentiment_summary

# Function to generate sentiment chart (pie chart)
def generate_sentiment_pie_chart(sentiment_summary):
    labels = sentiment_summary.keys()
    sizes = sentiment_summary.values()
    colors = ['green', 'gold', 'red']

    plt.figure(figsize=(6, 6))
    plt.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
    plt.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.

    chart_filename = 'sentiment_pie_chart.png'
    plt.savefig(chart_filename)
    plt.close()
    return chart_filename

# Function to create the PDF report for sentiment analysis
def create_sentiment_report_pdf(meeting_data):
    analysis, sentiment_summary = analyze_speech(meeting_data['transcriptData'])

    # Generate sentiment chart and get its filename
    chart_filename = generate_sentiment_pie_chart(sentiment_summary)

    file_name = f"{meeting_data['meetingTitle']}_sentiment_report.pdf"
    pdf = SimpleDocTemplate(file_name, pagesize=A4)
    elements = []
    
    styles = getSampleStyleSheet()
    title = "Sentiment Analysis Report"
    elements.append(Paragraph(title, styles['Title']))
    elements.append(Spacer(1, 12))

    # Add the pie chart to the PDF
    elements.append(Image(chart_filename, width=200, height=200))
    elements.append(Spacer(1, 12))

    for entry in analysis:
        elements.append(Paragraph(f"Speaker: {entry['speaker']}", styles['Normal']))
        elements.append(Paragraph(f"Sentiment: {entry['sentiment_category']}", styles['Normal']))
        elements.append(Paragraph(f"Content: {entry['content']}", styles['Normal']))
        elements.append(Spacer(1, 12))

    pdf.build(elements)
    print(f"{file_name} generated successfully.")

# Function to create the DOCX report for sentiment analysis
def create_sentiment_report_docx(meeting_data):
    analysis, sentiment_summary = analyze_speech(meeting_data['transcriptData'])

    # Generate sentiment chart and get its filename
    chart_filename = generate_sentiment_pie_chart(sentiment_summary)

    file_name = f"{meeting_data['meetingTitle']}_sentiment_report.docx"
    doc = Document()
    doc.add_heading("Sentiment Analysis Report", 0)

    # Add the sentiment pie chart image
    doc.add_picture(chart_filename, width=docx.shared.Inches(3), height=docx.shared.Inches(3))

    for entry in analysis:
        doc.add_paragraph(f"Speaker: {entry['speaker']}")
        doc.add_paragraph(f"Sentiment: {entry['sentiment_category']}")
        doc.add_paragraph(f"Content: {entry['content']}")
        doc.add_paragraph()  # Add a blank line

    doc.save(file_name)
    print(f"{file_name} generated successfully.")

def create_report_with_interval_sections_docx(meeting_data, interval_minutes):
    file_name = "Complete_Interval_Report.docx"
    doc = Document()

    doc.add_heading(f'Meeting Report: {meeting_data["meetingTitle"]}', level=1)

    # Ensure proper timestamp handling
    start_time = datetime.fromisoformat(meeting_data['meetingStartTimeStamp'].replace("Z", "+00:00"))
    end_time = datetime.fromisoformat(meeting_data['meetingEndTimeStamp'].replace("Z", "+00:00"))

    interval_seconds = interval_minutes * 60
    current_time = start_time

    # Process intervals and add sections to the report
    while current_time < end_time:
        interval_end_time = min(current_time + timedelta(seconds=interval_seconds), end_time)

        # Filter transcript data for the current interval
        interval_transcript_data = [
            entry for entry in meeting_data['transcriptData']
            if current_time <= datetime.fromisoformat(entry['timestamp'].replace("Z", "+00:00")) < interval_end_time
        ]

        if interval_transcript_data:
            # Create section title for the interval
            interval_title = f"Interval: {current_time.strftime('%H:%M')} - {interval_end_time.strftime('%H:%M')}"
            doc.add_heading(interval_title, level=2)

            # Analyze speech for this interval
            analysis_results, sentiment_summary = analyze_speech(interval_transcript_data)

            # Generate overall summary and key takeaways
            overall_summary, key_takeaways = generate_summary_and_takeaways(analysis_results)

            # Add overall summary for this interval
            doc.add_heading("Summary:", level=3)
            doc.add_paragraph(overall_summary)

            # Add key takeaways as bullet points
            doc.add_heading("Key Takeaways:", level=3)
            for takeaway in key_takeaways:
                doc.add_paragraph(f"• {takeaway}", style='ListBullet')

            # Add sentiment summary for this interval
            doc.add_heading("Sentiment Summary:", level=3)
            for sentiment, count in sentiment_summary.items():
                doc.add_paragraph(f"{sentiment}: {count} occurrences")

        # Move to the next interval
        current_time = interval_end_time

    # Save the DOCX file
    doc.save(file_name)
    print(f"{file_name} generated successfully.")

# Update main function to generate reports based on user input
def generate_reports(meeting_data):
    interval = int(input("Enter the time interval for sections (in minutes): "))
    print("Select the report type to generate:")
    print("1. Normal Report")
    print("2. Speaker Ranking Report")
    print("3. Sentiment Report")
    print("4. Interval Report")
    report_choice = input("Enter the number corresponding to your choice: ")

    print("Select the format:")
    print("1. PDF")
    print("2. DOCX")
    format_choice = input("Enter the number corresponding to your choice: ")

    if report_choice == '1':
        if format_choice == '1':
            create_normal_report_pdf(meeting_data)
        elif format_choice == '2':
            create_normal_report_docx(meeting_data)
    elif report_choice == '2':
        if format_choice == '1':
            create_speaker_ranking_report_pdf(meeting_data)
        elif format_choice == '2':
            create_speaker_ranking_report_docx(meeting_data)
    elif report_choice == '3':
        if format_choice == '1':
            create_sentiment_report_pdf(meeting_data)
        elif format_choice == '2':
            create_sentiment_report_docx(meeting_data)
    elif report_choice == '4':
        if format_choice == '1':
            create_report_with_interval_sections(meeting_data, interval)
        elif format_choice == '2':
            create_report_with_interval_sections_docx(meeting_data, interval)  # Corrected line
    else:
        print("Invalid choice.")

# Example meeting data (as before)
meeting_data = {
    'meetingTitle': 'Team Sync',
    'meetingStartTimeStamp': '2024-09-28T10:00:00.000Z',
    'meetingEndTimeStamp': '2024-09-28T11:30:00.000Z',
    'attendees': ['Alice', 'Bob', 'Charlie'],
    'transcriptData': [
        {'name': 'Alice', 'content': 'Hello everyone! How are you?', 'timestamp': '2024-09-28T10:05:00.000Z'},
        {'name': 'Bob', 'content': 'I am fine, thanks!', 'timestamp': '2024-09-28T10:10:00.000Z'},
        {'name': 'Charlie', 'content': 'Good morning, team!', 'timestamp': '2024-09-28T10:20:00.000Z'},
        {'name': 'Alice', 'content': 'Let’s discuss the project update.', 'timestamp': '2024-09-28T10:30:00.000Z'},
        {'name': 'Bob', 'content': 'We have completed the first phase.', 'timestamp': '2024-09-28T10:50:00.000Z'}
    ],
    'speakerDuration': {
        'Alice': 10,
        'Bob': 7,
        'Charlie': 5
    }
}

# Generate reports
generate_reports(meeting_data)
