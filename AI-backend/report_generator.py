from datetime import datetime, timedelta
import matplotlib # type: ignore
import os
import matplotlib.pyplot as plt # type: ignore
from reportlab.lib.pagesizes import A4 # type: ignore
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem,Image # type: ignore
from reportlab.lib.styles import getSampleStyleSheet # type: ignore # type: ignore
from textblob import TextBlob # type: ignore
from docx import Document # type: ignore
from datetime import datetime, timedelta
from reportlab.lib import colors # type: ignore
from reportlab.lib.pagesizes import letter # type: ignore
from reportlab.platypus import Table, TableStyle # type: ignore
import docx # type: ignore
from docx.shared import Inches # type: ignore
import re
from collections import Counter
from textblob import TextBlob # type: ignore
from transformers import pipeline # type: ignore
import torch # type: ignore
matplotlib.use('Agg')  # Use non-GUI backend

# Function to categorize sentiment based on polarity score
# Load the summarization model
summarizer = pipeline("summarization", model="facebook/bart-large-cnn",device=torch.device("mps") if torch.backends.mps.is_available() else torch.device("cpu"))

# Function to categorize sentiment based on polarity score
def categorize_sentiment(polarity):
    if polarity > 0.2:
        return "Positive"
    elif polarity < -0.2:
        return "Negative"
    else:
        return "Neutral"

# Function to analyze speech and categorize sentiment
def analyze_speech(transcript_data):
    """Analyze speech data to categorize sentiment."""
    analysis_results = []
    sentiment_summary = {"Positive": 0, "Neutral": 0, "Negative": 0}

    for entry in transcript_data:
        speaker = entry['name']
        content = entry['content']

        # Perform sentiment analysis
        blob = TextBlob(content)
        sentiment = blob.sentiment
        sentiment_category = categorize_sentiment(sentiment.polarity)

        # Increment sentiment summary count
        sentiment_summary[sentiment_category] += 1

        results = {
            'speaker': speaker,
            'content': content,
            'polarity': sentiment.polarity,
            'subjectivity': sentiment.subjectivity,
            'sentiment_category': sentiment_category
        }
        analysis_results.append(results)
    
    return analysis_results, sentiment_summary


def create_report_with_interval_sections_pdf(meeting_data, interval_minutes):
    """Generates a PDF report divided into time intervals (in minutes) with summaries and key takeaways."""
    
    file_name = "./reports/Complete_Interval_Report.pdf"
    pdf = SimpleDocTemplate(file_name, pagesize=A4)
    elements = []
    
    styles = getSampleStyleSheet()
    title = f"<b>Meeting Report: {meeting_data['meetingTitle']}</b>"
    elements.append(Paragraph(title, styles['Title']))
    elements.append(Spacer(1, 12))

    start_time = datetime.fromisoformat(meeting_data['meetingStartTimeStamp'].replace("Z", "+00:00"))
    end_time = datetime.fromisoformat(meeting_data['meetingEndTimeStamp'].replace("Z", "+00:00"))
    
    interval_seconds = interval_minutes * 60
    current_time = start_time

    # Process intervals and add sections to the report
    while current_time < end_time:
        interval_end_time = min(current_time + timedelta(seconds=interval_seconds), end_time)

        # Filter transcript data for the current interval
        interval_transcript_data = [
            entry for entry in meeting_data['transcriptData']
            if current_time <= datetime.fromisoformat(entry['timestamp'].replace("Z", "+00:00")) < interval_end_time
        ]
        
        if interval_transcript_data:
            # Create section title for the interval
            interval_title = f"<b>Interval: {current_time.strftime('%H:%M')} - {interval_end_time.strftime('%H:%M')}</b>"
            elements.append(Paragraph(interval_title, styles['Heading2']))
            elements.append(Spacer(1, 12))
            
            # Analyze speech for this interval
            analysis_results, sentiment_summary = analyze_speech(interval_transcript_data)

            # Generate overall summary and key takeaways
            overall_summary, key_takeaways = generate_summary_and_takeaways(analysis_results)

            # Add overall summary for this interval
            elements.append(Paragraph("<b>Summary:</b>", styles['Heading3']))
            elements.append(Paragraph(overall_summary, styles['Normal']))
            elements.append(Spacer(1, 12))

            # Add key takeaways as paragraphs
            elements.append(Paragraph("<b>Key Takeaways:</b>", styles['Heading3']))
            for takeaway in key_takeaways.split('.'):
                if takeaway.strip():
                    elements.append(Paragraph(takeaway.strip(), styles['Normal']))
                    elements.append(Spacer(1, 6))  # Space between takeaways

            # Add sentiment summary for this interval
            elements.append(Paragraph("<b>Sentiment Summary:</b>", styles['Heading3']))
            for sentiment, count in sentiment_summary.items():
                elements.append(Paragraph(f"{sentiment}: {count} occurrences", styles['Normal']))
                elements.append(Spacer(1, 12))
        
        current_time = interval_end_time

    # Build the final PDF
    pdf.build(elements)
    print(f"{file_name} generated successfully.")
    return file_name

def create_report_with_interval_sections_docx(meeting_data, interval_minutes):
    """Generates a DOCX report divided into time intervals (in minutes) with summaries and key takeaways."""
    
    file_name = "./reports/Complete_Interval_Report.docx"
    doc = Document()

    # Add title
    doc.add_heading(f'Meeting Report: {meeting_data["meetingTitle"]}', level=1)

    # Ensure proper timestamp handling
    start_time = datetime.fromisoformat(meeting_data['meetingStartTimeStamp'].replace("Z", "+00:00"))
    end_time = datetime.fromisoformat(meeting_data['meetingEndTimeStamp'].replace("Z", "+00:00"))

    current_time = start_time

    # Process intervals and add sections to the report
    while current_time < end_time:
        interval_end_time = min(current_time + timedelta(seconds=interval_minutes * 60), end_time)

        # Filter transcript data for the current interval
        interval_transcript_data = [
            entry for entry in meeting_data['transcriptData']
            if current_time <= datetime.fromisoformat(entry['timestamp'].replace("Z", "+00:00")) < interval_end_time
        ]

        if interval_transcript_data:
            # Create section title for the interval
            interval_title = f"Interval: {current_time.strftime('%H:%M')} - {interval_end_time.strftime('%H:%M')}"
            doc.add_heading(interval_title, level=2)

            # Analyze speech for this interval
            analysis_results, sentiment_summary = analyze_speech(interval_transcript_data)

            # Generate overall summary and key takeaways
            overall_summary, key_takeaways = generate_summary_and_takeaways(analysis_results)

            # Add overall summary for this interval
            doc.add_heading("Summary:", level=3)
            doc.add_paragraph(overall_summary)

            # Add key takeaways as paragraphs
            doc.add_heading("Key Takeaways:", level=3)
            for takeaway in key_takeaways.split('.'):
                if takeaway.strip():
                    doc.add_paragraph(takeaway.strip())

        # Move to the next interval
        current_time = interval_end_time

    # Save the DOCX file
    doc.save(file_name)
    print(f"{file_name} generated successfully.")
    return file_name

# Function to dynamically generate the overall summary based on the transcript
import re
from reportlab.lib.pagesizes import A4 # type: ignore # type: ignore
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem # type: ignore

from collections import Counter
from transformers import pipeline # type: ignore
from reportlab.lib.pagesizes import A4 # type: ignore
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem # type: ignore

def summarize_text(text, max_length=50, min_length=5):
    """Summarize the provided text."""
    if len(text.strip()) == 0:
        return ""
    
    input_length = len(text.split())
    max_length = min(max_length, input_length)  # Adjust max_length if input is shorter

    summary = summarizer(text, max_length=max_length, min_length=min_length, do_sample=False)
    return summary[0]['summary_text']

def summarize_takeaways(takeaways):
    """Summarize key takeaways using the summarization model."""
    combined_takeaways = " ".join(takeaways)
    summary = summarize_text(combined_takeaways)
    return summary

def is_meaningful(content):
    """Check if the content is meaningful (not too short and no question mark)."""
    return len(content.split()) >= 3 and '?' not in content

def generate_summary_and_takeaways(analysis_results):
    """Generate overall summary and key takeaways based on analysis."""
    overall_summary = "Summary of the meeting:"
    key_takeaways = []

    if analysis_results:
        for entry in analysis_results:
            speaker = entry['speaker']
            sentiment = entry['sentiment_category']
            content = entry['content']
            
            # Improved phrasing for clarity
            takeaway = f"{speaker} expressed a {sentiment.lower()} sentiment, stating: '{content}'"
            key_takeaways.append(takeaway)

    # Summarize key takeaways using the summarization model
    summarized_takeaways = summarize_takeaways(key_takeaways)
    
    return overall_summary, summarized_takeaways

# Function to generate the overall summary of the meeting
def generate_overall_summary(transcript_data):
    content_list = [entry['content'].strip() for entry in transcript_data if is_meaningful(entry['content'])]
    summary_text = " ".join(content_list)  # Combine content for better context

    # Summarize the overall content
    summarized_summary = summarize_text(summary_text, max_length=150, min_length=30)
    return f"In summary, the meeting covered the following key points: {summarized_summary}"

# Improved function to generate key takeaways
def generate_key_takeaways(transcript_data):
    takeaways = []

    for entry in transcript_data:
        content = entry['content'].strip()
        if is_meaningful(content):
            takeaways.append(f"{entry['name']} mentioned: {content}")


    if not takeaways:
        takeaways.append("There were no specific key takeaways from the meeting.")

    # Combine and summarize key takeaways for better context
    combined_takeaways = " ".join(takeaways)
    summarized_takeaways = summarize_text(combined_takeaways, max_length=150, min_length=30)
    
    return summarized_takeaways

# Function to generate speaker summaries
def generate_speaker_summaries(transcript_data, speaker_durations):
    speaker_summaries = {}

    for entry in transcript_data:
        speaker = entry['name']
        content = entry['content'].strip()

        if is_meaningful(content):
            if speaker not in speaker_summaries:
                speaker_summaries[speaker] = {
                    "points": [],
                    "duration": speaker_durations.get(speaker, 0)
                }
            speaker_summaries[speaker]['points'].append(content)

    # Summarize each speaker's contributions
    for speaker, data in speaker_summaries.items():
        combined_points = " ".join(data['points'])
        summarized_points = summarize_text(combined_points, max_length=150, min_length=30)
        speaker_summaries[speaker]['summary'] = f"{speaker} contributed: {summarized_points}"

    return speaker_summaries

def format_time(timestamp):
    """Convert ISO format timestamp to a readable format with AM/PM."""
    dt = datetime.fromisoformat(timestamp.replace("Z", "+00:00"))
    return dt.strftime("%Y-%m-%d %I:%M %p")  

from reportlab.platypus import Image # type: ignore
from reportlab.lib.units import inch # type: ignore

def create_normal_report_pdf(meeting_data):
    file_name = f"./reports/{meeting_data['meetingTitle']}_summary_report.pdf"
    pdf = SimpleDocTemplate(file_name, pagesize=A4)
    elements = []

    styles = getSampleStyleSheet()
    title = f"<b>{meeting_data['meetingTitle']}</b>"
    elements.append(Paragraph(title, styles['Title']))
    elements.append(Spacer(1, 12))

    # Add convenor and meeting details
    convenor_info = f"Convenor: {meeting_data['convenor']}<br/>"
    meeting_details = (
        convenor_info +
        f"Meeting Start Time: {format_time(meeting_data['meetingStartTimeStamp'])}<br/>"
        f"Meeting End Time: {format_time(meeting_data['meetingEndTimeStamp'])}<br/>"
        f"Attendees: {', '.join(meeting_data.get('attendees', ['N/A']))}<br/>"
        f"Speakers: {', '.join(meeting_data.get('speakers', ['N/A']))}<br/><br/>"
    )
    elements.append(Paragraph(meeting_details, styles['Normal']))

    # Generate and add Executive Summary
    overall_summary = generate_overall_summary(meeting_data.get('transcriptData', []))
    elements.append(Paragraph("<b>Executive Summary:</b>", styles['Heading2']))
    elements.append(Paragraph(overall_summary, styles['Normal']))
    elements.append(Spacer(1, 12))

    # Generate and add Key Takeaways
    key_takeaways = generate_key_takeaways(meeting_data.get('transcriptData', []))
    elements.append(Paragraph("<b>Key Takeaways:</b>", styles['Heading2']))
    elements.append(Paragraph(key_takeaways, styles['Normal']))
    elements.append(Spacer(1, 12))

    # Generate and add Speaker Summaries
    speaker_summaries = generate_speaker_summaries(meeting_data.get('transcriptData', []), meeting_data.get('speakerDuration', {}))
    elements.append(Paragraph("<b>Speaker Summaries:</b>", styles['Heading2']))
    for speaker, data in speaker_summaries.items():
        speaker_info = f"<b>{speaker}</b><br/>Total Speaking Time: {data['duration']} seconds<br/>"
        elements.append(Paragraph(speaker_info, styles['Normal']))
        elements.append(Paragraph(data['summary'], styles['Normal']))  # Show summarized points
        elements.append(Spacer(1, 12))

    # Add screenshots section
    if 'screenshots' in meeting_data and meeting_data['screenshots']:
        elements.append(Paragraph("<b>Screenshots:</b>", styles['Heading2']))
        for screenshot in meeting_data['screenshots']:
            elements.append(Paragraph(f"Screenshot taken by {screenshot['takenBy']} at {format_time(screenshot['timestamp'])}", styles['Normal']))
            try:
                image_path = screenshot['takenBy'] + '/' + screenshot['filename']
                img = Image('../node_backend/screenshots/'+ image_path, width=4*inch, height=3*inch)  # Adjust the size as needed
                elements.append(img)
                elements.append(Spacer(1, 12))
            except Exception as e:
                elements.append(Paragraph(f"Could not load image: {image_path}. Error: {str(e)}", styles['Normal']))
                elements.append(Spacer(1, 12))

    # Finalize and build the PDF
    pdf.build(elements)
    print(f"PDF report '{file_name}' generated successfully.")
    return file_name


# Function to create the normal report DOCX
def create_normal_report_docx(meeting_data):
    file_name = f"./reports/{meeting_data['meetingTitle']}_summary_report.docx"
    doc = Document()

    doc.add_heading(meeting_data['meetingTitle'], level=1)

    # Add convenor and meeting details
    convenor_info = f"Convenor: {meeting_data['convenor']}\n"
    doc.add_paragraph(convenor_info)
    doc.add_paragraph(
        f"Meeting Start Time: {format_time(meeting_data['meetingStartTimeStamp'])}\n"
        f"Meeting End Time: {format_time(meeting_data['meetingEndTimeStamp'])}\n"
        f"Attendees: {', '.join(meeting_data.get('attendees', ['N/A']))}\n"
        f"Speakers: {', '.join(meeting_data.get('speakers', ['N/A']))}\n"  # Added speakers
    )

    # Generate and add Executive Summary
    overall_summary = generate_overall_summary(meeting_data.get('transcriptData', []))
    doc.add_heading('Executive Summary:', level=2)
    doc.add_paragraph(overall_summary)

    # Generate and add Key Takeaways
    key_takeaways = generate_key_takeaways(meeting_data.get('transcriptData', []))
    doc.add_heading('Key Takeaways:', level=2)
    doc.add_paragraph(key_takeaways)

    # Generate and add Speaker Summaries
    speaker_summaries = generate_speaker_summaries(meeting_data.get('transcriptData', []), meeting_data.get('speakerDuration', {}))
    doc.add_heading('Speaker Summaries:', level=2)
    for speaker, data in speaker_summaries.items():
        doc.add_paragraph(f"{speaker} - Total Speaking Time: {data['duration']} seconds")
        for point in data['points']:
            doc.add_paragraph(f"• {point}")

    # Add screenshots section
    if 'screenshots' in meeting_data and meeting_data['screenshots']:
        doc.add_heading('Screenshots:', level=2)
        for screenshot in meeting_data['screenshots']:
            doc.add_paragraph(f"Screenshot taken by {screenshot['takenBy']} at {format_time(screenshot['timestamp'])}")
            try:
                image_path ='../node_backend/screenshots/' + screenshot['takenBy'] + '/' + screenshot['filename']
                img = Image(image_path, width=4*inch, height=3*inch)
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(4), height=Inches(3))
                    doc.add_paragraph()  # Adds an empty paragraph to create space between images
                else:
                    doc.add_paragraph(f"Could not load image: {image_path}. File not found.")
            except Exception as e:
                doc.add_paragraph(f"Could not load image: {image_path}. Error: {str(e)}")


    # Save the DOCX file
    doc.save(file_name)
    print(f"DOCX report '{file_name}' generated successfully.")
    return file_name

def create_speaker_ranking_report_pdf(meeting_data):
    file_name = "./reports/Speaker_Ranking_Report.pdf"
    pdf = SimpleDocTemplate(file_name, pagesize=A4)
    elements = []
    
    styles = getSampleStyleSheet()
    title = f"<b>Speaker Ranking Report: {meeting_data['meetingTitle']}</b>"
    elements.append(Paragraph(title, styles['Title']))
    elements.append(Spacer(1, 12))

    # Prepare to collect speaker data
    speaker_data = {}

    # Process transcript data to gather speaker information
    for entry in meeting_data['transcriptData']:
        speaker_name = entry['name']
        content = entry['content']

        if speaker_name not in speaker_data:
            speaker_data[speaker_name] = {
                'count': 0,
                'summary': []  # List to hold the speaker's speech content
            }
        
        speaker_data[speaker_name]['count'] += 1
        speaker_data[speaker_name]['summary'].append(content)  # Collect each speech content

    # Sort speakers by their count (descending)
    sorted_speakers = sorted(speaker_data.items(), key=lambda item: item[1]['count'], reverse=True)

    # Add speaker ranking to the PDF
    elements.append(Paragraph("<b>Speaker Ranking:</b>", styles['Heading2']))
    for rank, (speaker_name, data) in enumerate(sorted_speakers, start=1):
        elements.append(Paragraph(f"{rank}. {speaker_name} - {data['count']} contributions", styles['Normal']))
        
        # Summarize the speaker's contributions using existing summarize_text function
        combined_summary = " ".join(data['summary'])
        summarized_contributions = summarize_text(combined_summary)

        # Add summarized contributions for each speaker
        elements.append(Paragraph("<b>Summary of Contributions:</b>", styles['Heading3']))
        elements.append(Paragraph(summarized_contributions, styles['Normal']))
        elements.append(Spacer(1, 6))  # Space between summaries

    # Build the final PDF
    pdf.build(elements)
    print(f"{file_name} generated successfully.")
    return file_name

def create_speaker_ranking_report_docx(meeting_data):
    file_name = "./reports/Speaker_Ranking_Report.docx"
    doc = Document()

    doc.add_heading(f'Speaker Ranking Report: {meeting_data["meetingTitle"]}', level=1)

    # Prepare to collect speaker data
    speaker_data = {}

    # Process transcript data to gather speaker information
    for entry in meeting_data['transcriptData']:
        speaker_name = entry['name']
        content = entry['content']
        
        if speaker_name not in speaker_data:
            speaker_data[speaker_name] = {
                'count': 0,
                'summary': []  # List to hold the speaker's speech content
            }
        
        speaker_data[speaker_name]['count'] += 1
        speaker_data[speaker_name]['summary'].append(content)  # Collect each speech content

    # Sort speakers by their count (descending)
    sorted_speakers = sorted(speaker_data.items(), key=lambda item: item[1]['count'], reverse=True)

    # Add speaker ranking to the document
    doc.add_heading('Speaker Ranking:', level=2)
    for rank, (speaker_name, data) in enumerate(sorted_speakers, start=1):
        doc.add_paragraph(f"{rank}. {speaker_name} - {data['count']} contributions")

        # Summarize the speaker's contributions using existing summarize_text function
        combined_summary = " ".join(data['summary'])
        summarized_contributions = summarize_text(combined_summary)

        # Add summarized contributions for each speaker
        doc.add_heading("Summary of Contributions:", level=3)
        doc.add_paragraph(summarized_contributions)

    # Save the DOCX file
    doc.save(file_name)
    print(f"{file_name} generated successfully.")
    return file_name

# Function to categorize sentiment based on polarity score
def categorize_sentiment(polarity):
    if polarity > 0.2:
        return "Positive"
    elif polarity < -0.2:
        return "Negative"
    else:
        return "Neutral"

# Function to analyze speech and categorize sentiment
def analyze_speech(transcript_data):
    analysis_results = []
    sentiment_summary = {"Positive": 0, "Neutral": 0, "Negative": 0}

    for entry in transcript_data:
        speaker = entry['name']
        content = entry['content']

        # Perform sentiment analysis
        blob = TextBlob(content)
        sentiment = blob.sentiment
        sentiment_category = categorize_sentiment(sentiment.polarity)

        # Increment sentiment summary count
        sentiment_summary[sentiment_category] += 1

        results = {
            'speaker': speaker,
            'content': content,
            'polarity': sentiment.polarity,
            'subjectivity': sentiment.subjectivity,
            'sentiment_category': sentiment_category
        }
        analysis_results.append(results)
    
    return analysis_results, sentiment_summary

# Function to generate sentiment chart (pie chart)
def generate_sentiment_pie_chart(sentiment_summary):
    labels = sentiment_summary.keys()
    sizes = sentiment_summary.values()
    colors = ['green', 'gold', 'red']

    plt.figure(figsize=(6, 6))
    plt.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
    plt.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.

    chart_filename = './reports/sentiment_pie_chart.png'
    plt.savefig(chart_filename)
    plt.close()
    return chart_filename

# Function to create the PDF report for sentiment analysis
def create_sentiment_report_pdf(meeting_data):
    analysis, sentiment_summary = analyze_speech(meeting_data['transcriptData'])

    # Generate sentiment chart and get its filename
    chart_filename = generate_sentiment_pie_chart(sentiment_summary)

    file_name = f"./reports/{meeting_data['meetingTitle']}_sentiment_report.pdf"
    pdf = SimpleDocTemplate(file_name, pagesize=A4)
    elements = []
    
    styles = getSampleStyleSheet()
    title = "Sentiment Analysis Report"
    elements.append(Paragraph(title, styles['Title']))
    elements.append(Spacer(1, 12))

    # Add the pie chart to the PDF
    elements.append(Image(chart_filename, width=200, height=200))
    elements.append(Spacer(1, 12))

    for entry in analysis:
        elements.append(Paragraph(f"Speaker: {entry['speaker']}", styles['Normal']))
        elements.append(Paragraph(f"Sentiment: {entry['sentiment_category']}", styles['Normal']))
        elements.append(Paragraph(f"Content: {entry['content']}", styles['Normal']))
        elements.append(Spacer(1, 12))

    pdf.build(elements)
    print(f"{file_name} generated successfully.")
    return file_name

# Function to create the DOCX report for sentiment analysis
def create_sentiment_report_docx(meeting_data):
    analysis, sentiment_summary = analyze_speech(meeting_data['transcriptData'])

    # Generate sentiment chart and get its filename
    chart_filename = generate_sentiment_pie_chart(sentiment_summary)

    file_name = f"./reports/{meeting_data['meetingTitle']}_sentiment_report.docx"
    doc = Document()
    doc.add_heading("Sentiment Analysis Report", 0)

    # Add the sentiment pie chart image
    doc.add_picture(chart_filename, width=docx.shared.Inches(3), height=docx.shared.Inches(3))

    for entry in analysis:
        doc.add_paragraph(f"Speaker: {entry['speaker']}")
        doc.add_paragraph(f"Sentiment: {entry['sentiment_category']}")
        doc.add_paragraph(f"Content: {entry['content']}")
        doc.add_paragraph()  # Add a blank line

    doc.save(file_name)
    print(f"{file_name} generated successfully.")
    return file_name

class NormalReport:
    def __init__(self):
        pass
    def __str__(self):
        return "Normal Report"

class SpeakerRankingReport:
    def __init__(self):
        pass
    def __str__(self):
        return "Speaker Ranking Report"

class SentimentReport:
    def __init__(self):
        pass
    def __str__(self):
        return "Sentiment Report"
    
class IntervalReport:
    def __init__(self):
        pass
    def __str__(self):
        return "Interval Report"
    
class PDF_Type:
    def __init__(self):
        pass
    def __str__(self):
        return "PDF"

class DOCX_Type:
    def __init__(self):
        pass
    def __str__(self):
        return "DOCX"
    
def create_reports_directory():
    '''Creates a directory named './reports' if it does not exist.'''
    import os
    if not os.path.exists('./reports'):
        os.makedirs('./reports')
        print("Reports directory created successfully.")

# Main function to generate reports based on user input
def generate_reports(meeting_data,report_choice=NormalReport,format_choice=PDF_Type,interval_minutes=5):
    '''Generates and saves a report based on user input.
        meeting_data: dict, contains meeting details and transcript data
        report_choice: the type of report to generate. Use the classes from this file: NormalReport, SpeakerRankingReport, SentimentReport, IntervalReport
        format_choice: the format of the report. Use the classes from this file: PDF, DOCX'''
    
    # Create reports directory if it does not exist
    create_reports_directory()


    if report_choice ==  NormalReport:
        if format_choice ==  PDF_Type:
            return create_normal_report_pdf(meeting_data)
        elif format_choice ==  DOCX_Type:
            return create_normal_report_docx(meeting_data)
    elif report_choice ==  SpeakerRankingReport:
        print("generating speaker ranking report")
        if format_choice ==  PDF_Type:
            return create_speaker_ranking_report_pdf(meeting_data)
        elif format_choice ==  DOCX_Type:
            return create_speaker_ranking_report_docx(meeting_data)
    elif report_choice ==  SentimentReport:
        if format_choice ==  PDF_Type:
            return create_sentiment_report_pdf(meeting_data)
        elif format_choice ==  DOCX_Type:
            return create_sentiment_report_docx(meeting_data)
    elif report_choice ==  IntervalReport:
        if not interval_minutes:
            raise ValueError("Interval minutes must be provided for Interval Report.")
        if format_choice ==  PDF_Type:
            return create_report_with_interval_sections_pdf(meeting_data, interval_minutes)
        elif format_choice ==  DOCX_Type:
            return create_report_with_interval_sections_docx(meeting_data, interval_minutes)
    else:
        raise ValueError("Invalid report or format choice.")

# Example usage
if __name__ == "__main__":
    # Sample meeting data object for testing
    meeting_data = {    
        'meetingTitle': 'yzf-wdcr-zam',
        'convenor': 'TitaNyte Official',  
        'speakers': ['TitaNyte Official', 'Prateek'],  
        'meetingStartTimeStamp': '2024-09-29T12:20:48.000Z',
        'meetingEndTimeStamp': '2024-09-29T12:26:09.000Z',
        'attendees': ['Prateek', 'TitaNyte Official'],
        'transcriptData': [
            {'name': 'TitaNyte Official', 'content': 'Hi, how are you?', 'timestamp': '2024-09-29T12:21:37.000Z'},
            {'name': 'Prateek', 'content': "Hi, I'm fine. So what's going to be the agenda of today's discussion", 'timestamp': '2024-09-29T12:21:38.000Z'},
            {'name': 'TitaNyte Official', 'content': 'So now, we\'ll be focusing.', 'timestamp': '2024-09-29T12:21:43.000Z'},
            {'name': 'Prateek', 'content': 'I say, Let\'s discuss about the features that are extension, is going to give to the users, okay? So the first and foremost, is that it will give the transcriptions, Which it will basically. Web script actually it will just scrape it from the captions.', 'timestamp': '2024-09-29T12:21:45.000Z'},
            {'name': 'Prateek', 'content': 'Google Meet is generating. And what will happen afterwards. Once that transcripts are there in the database, can you please tell me?', 'timestamp': '2024-09-29T12:22:03.000Z'},
            {'name': 'TitaNyte Official', 'content': 'Sure. so once US transcripts are stored in the database, They?', 'timestamp': '2024-09-29T12:22:12.000Z'},
            {'name': 'Prateek', 'content': "It's okay, it's okay. You can continue", 'timestamp': '2024-09-29T12:22:21.000Z'},
            {'name': 'TitaNyte Official', 'content': 'Okay, so once the transit are stored So once the transit are stored in the database, what you\'re doing is we then have an ML model or basically AI model, right? A generative AI model which basically takes these transcripts and then tries to create reports BAS', 'timestamp': '2024-09-29T12:22:24.000Z'},
            {'name': 'TitaNyte Official', 'content': 'filters that we are working on, right? So the filters that we so far, basically, it includes speaker based reporting, right? It includes interval-based reporting. And a host of other features. so, this training part happens when you click on Generate,', 'timestamp': '2024-09-29T12:22:40.000Z'},
            {'name': 'TitaNyte Official', 'content': 'front end, right? So you have an extension, you have a front end and on the front end you log in using Google, Right? So this is how it works. And so as soon as you click on Generate button, there\'ll be a model that pop up, right? And on the model, you', 'timestamp': '2024-09-29T12:23:02.000Z'},
            {'name': 'TitaNyte Official', 'content': 'You can enter the meeting title, you can enter the the kind of report that you want whether you want to talk X format or P or PDF format right and basically the format stuff, right? So with formatting want and then the title and then the option and then', 'timestamp': '2024-09-29T12:23:22.000Z'},
            {'name': 'TitaNyte Official', 'content': 'add Prompts as well, which would enhance the report creation. So that is something that\'s on the cards. Apart from that in the extension, you can also take screenshots, right? Which makes it easy to keep track of what had been shared during the meeting.', 'timestamp': '2024-09-29T12:23:42.000Z'},
            {'name': 'TitaNyte Official', 'content': 'So let me try taking a screenshot now. Let\'s see if it works, and let\'s see if it works here. I think we\'ve grabbed a screenshot we need to check we\'ll check that later. So yeah that\'s how we take a So yeah that\'s how we take a screenshot basically you clear?', 'timestamp': '2024-09-29T12:23:59.000Z'},
            {'name': 'TitaNyte Official', 'content': 'Basically you go to the Extension, you click on the extension, and then you click on the Capture button there. So nice work, I like the user interface that you\'ve created for the extension.', 'timestamp': '2024-09-29T12:24:17.000Z'},
            {'name': 'Prateek', 'content': 'Okay, so now I will be covering like how anyone can use our extension. Like, by through our repo First, they\'ll have to clone the repository. Then they will have to install all the packages. And after I know, it is already available on the Chrome store.', 'timestamp': '2024-09-29T12:24:30.000Z'},
            {'name': 'Prateek', 'content': 'still, if they want to run it on there, let\'s say, that\'s if they want to run it on their own device, then they will have to clone that apple. This is what I\'m going to add in the Readme, so please correct me if I\'m So please correct me if I\'m going wrong. Some', 'timestamp': '2024-09-29T12:24:50.000Z'},
            {'name': 'TitaNyte Official', 'content': 'Okay.', 'timestamp': '2024-09-29T12:25:01.000Z'},
            {'name': 'Prateek', 'content': 'And they will have to do the NPM install for both front end and backend and PIP install for the AI backend. Okay, and afterwards, they will have to set the environment variables, which we, which we will give it to them, just until the judgment has been', 'timestamp': '2024-09-29T12:25:02.000Z'},
            {'name': 'Prateek', 'content': 'after then we\'ll remove that remove it. So yeah that\'s it after then they will be just able to run it. run.', 'timestamp': '2024-09-29T12:25:17.000Z'},
            {'name': 'TitaNyte Official', 'content': 'Okay.', 'timestamp': '2024-09-29T12:25:27.000Z'},
            {'name': 'Prateek', 'content': 'What is it? Is it is there something which I\'m missing?', 'timestamp': '2024-09-29T12:25:27.000Z'},
            {'name': 'TitaNyte Official', 'content': 'Yeah, I think the few more things, I mean, you also have to set up the Python directory that', 'timestamp': '2024-09-29T12:25:31.000Z'},
            {'name': 'Prateek', 'content': 'Okay.', 'timestamp': '2024-09-29T12:25:37.000Z'},
            {'name': 'TitaNyte Official', 'content': 'add all the step instructions. I think you mentioned that, right?', 'timestamp': '2024-09-29T12:25:42.000Z'},
            {'name': 'Prateek', 'content': 'Yeah.', 'timestamp': '2024-09-29T12:25:41.000Z'},
            {'name': 'TitaNyte Official', 'content': 'Yeah, but then we also need to provide requirements or TC, which would be there. But apart from that, you also have to, you also have to add dot Env. Dot Emmy file, basically an example file, which shows in the format. shows in Right. So I think that\'s', 'timestamp': '2024-09-29T12:25:42.000Z'},
            {'name': 'TitaNyte Official', 'content': 'I think we\'re good to go, right?', 'timestamp': '2024-09-29T12:25:59.000Z'},
            {'name': 'Prateek', 'content': 'Yeah, we are good to go.', 'timestamp': '2024-09-29T12:26:02.000Z'},
            {'name': 'TitaNyte Official', 'content': 'so, I think', 'timestamp': '2024-09-29T12:26:03.000Z'},
            {'name': 'Prateek', 'content': 'Thank you.', 'timestamp': '2024-09-29T12:26:04.000Z'},
            {'name': 'TitaNyte Official', 'content': 'Yeah, thank you have', 'timestamp': '2024-09-29T12:26:05.000Z'},
            {'name': 'Prateek', 'content': 'Yeah.', 'timestamp': '2024-09-29T12:26:06.000Z'},
            {'name': 'TitaNyte Official', 'content': 'Thank you have a nice day.', 'timestamp': '2024-09-29T12:26:06.000Z'}
        ],
        'speakerDuration': {
            'TitaNyte Official': 163,
            'Prateek': 104
        },
        'screenshots': [
            {
                'filename': 'screenshot_1727619945027.png',
                'timestamp': '2024-09-29T12:24:05.299Z',
                'takenBy': 'afnanind@gmail.com'
            }
        ]
    }


    print("Returns",generate_reports(meeting_data, report_choice=NormalReport, format_choice=PDF_Type))
    print("Returns",generate_reports(meeting_data, report_choice=NormalReport, format_choice=DOCX_Type))
    print("Returns",generate_reports(meeting_data, report_choice=SpeakerRankingReport, format_choice=PDF_Type))
    print("Returns",generate_reports(meeting_data, report_choice=SpeakerRankingReport, format_choice=DOCX_Type))
    print("Returns",generate_reports(meeting_data, report_choice=SentimentReport, format_choice=PDF_Type))
    print("Returns",generate_reports(meeting_data, report_choice=SentimentReport, format_choice=DOCX_Type))
    print("Returns",generate_reports(meeting_data, report_choice=IntervalReport, format_choice=PDF_Type, interval_minutes=5))
    print("Returns",generate_reports(meeting_data, report_choice=IntervalReport, format_choice=DOCX_Type, interval_minutes=5))
    # Expected output: PDF and DOCX reports generated successfully in the 'reports' directory
